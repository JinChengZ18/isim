from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DOCX_PATH = Path(__file__).with_name("chapter03.docx")
BODY_WIDTH = Inches(6.5)


def set_east_asia_font(style, font_name):
    run_props = style._element.get_or_add_rPr()
    run_fonts = run_props.rFonts
    if run_fonts is None:
        run_fonts = OxmlElement("w:rFonts")
        run_props.insert(0, run_fonts)
    run_fonts.set(qn("w:eastAsia"), font_name)
    run_fonts.set(qn("w:hint"), "eastAsia")


def style_document(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    set_east_asia_font(normal, "宋体")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, 16, 8),
        "Heading 2": (14, 14, 6),
        "Heading 3": (12, 10, 4),
        "Heading 4": (11, 8, 4),
    }
    for name, (size, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        set_east_asia_font(style, "黑体")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.color.rgb = RGBColor(0, 0, 0)
    set_east_asia_font(caption, "宋体")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("图3-") or text.startswith("表3-"):
            paragraph.style = caption
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = text.startswith("表3-")
        elif text.startswith("注："):
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                run.font.size = Pt(9.5)

    for table in doc.tables:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        table.autofit = True

    for shape in doc.inline_shapes:
        if shape.width > BODY_WIDTH:
            scale = BODY_WIDTH / shape.width
            shape.width = int(shape.width * scale)
            shape.height = int(shape.height * scale)


def main():
    doc = Document(DOCX_PATH)
    style_document(doc)
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    main()
